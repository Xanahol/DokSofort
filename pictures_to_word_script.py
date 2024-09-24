import os
from docx import Document
from docx.shared import Inches
 
# Pfad zum Ordner mit den Bildern (bitte anpassen)
image_folder = r'C:\Users\Noël\siworks Dropbox\Noël Brand\PC\Documents\Scripts\Infra\20241909 Abnahme\20241909 Abnahme'
 
# Erstelle ein neues Word-Dokument
doc = Document()
 
# Durchlaufe alle Dateien im angegebenen Ordner
for filename in os.listdir(image_folder):
    if filename.endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif")):  # Erlaubte Bildformate
        # Füge den Dateinamen als Überschrift hinzu (ohne Dateiendung)
        doc.add_paragraph(filename.rsplit('.', 1)[0], style='Heading 2')
        # Füge das Bild ein
        image_path = os.path.join(image_folder, filename)
        doc.add_picture(image_path, width=Inches(5))  # Hier kannst du die Bildgröße anpassen
        # Füge eine Leerzeile hinzu
        doc.add_paragraph()
 
# Speichere das Dokument
doc.save(r'C:\Users\Noël\siworks Dropbox\Noël Brand\PC\Documents\Scripts\Infra\20241909 Abnahme\20241909 Abnahme\Bilderdokument.docx')
 
print("Dokument erstellt.")