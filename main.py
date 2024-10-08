import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docx.shared import Inches
from datetime import datetime
import threading
from PIL import Image, ImageTk

MAXIMUM_IMAGE_DIRS = 5

# Function to get the right path when running from PyInstaller


def resource_path(relative_path):
    """ Get the absolute path to a resource, works for dev and PyInstaller """
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)


def create_document(image_folders, save_path, progress_callback):
    doc = Document()
    total_images = sum(len(os.listdir(folder))
                       for folder in image_folders if os.path.isdir(folder))

    image_count = 0
    for image_folder in image_folders:
        for filename in os.listdir(image_folder):
            if filename.endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif")):
                # Add filename without extension as heading
                doc.add_paragraph(filename.rsplit(
                    '.', 1)[0], style='Heading 2')
                image_path = os.path.join(image_folder, filename)
                # Add the picture
                doc.add_picture(image_path, width=Inches(5))
                doc.add_paragraph()

                image_count += 1
                # Update the progress bar
                progress_callback(image_count, total_images)

    # Generate the current timestamp in the MMDDHHSS format
    timestamp = datetime.now().strftime("%m%d%H%M%S")

    # Save the document
    docx_file = os.path.join(save_path, f"DokSofort{timestamp}.docx")
    doc.save(docx_file)

    # Open the output folder
    os.startfile(save_path)

    messagebox.showinfo("Success", f"Document saved at: {save_path}")


# Function to select folders with images
def select_image_folder(directory_label, folder_list, add_button, generate_button, save_path_var, index):
    folder = filedialog.askdirectory(title="Select Folder with Images")
    if folder:
        folder_list[index] = folder
        # Show only the last characters
        directory_label.config(text=f"...{folder[-25:]}")
        check_generate_button_state(
            folder_list, save_path_var.get(), generate_button)


# Function to select where to save the document
def select_save_location(directory_label, save_path_var, folder_list, generate_button):
    folder = filedialog.askdirectory(title="Select Save Folder")
    if folder:
        save_path_var.set(folder)
        # Show only the last characters
        directory_label.config(text=f"...{folder[-25:]}")
        check_generate_button_state(
            folder_list, save_path_var.get(), generate_button)


# Check if Generate button should be enabled
def check_generate_button_state(folder_list, save_path, generate_button):
    if any(folder_list) and save_path:
        generate_button.config(state=tk.NORMAL)
    else:
        generate_button.config(state=tk.DISABLED)


# Function to show a progress bar while document is being generated
def show_progress_bar(total_images):
    progress_window = tk.Toplevel()
    progress_window.title("Generating Document...")
    progress_window.geometry("300x120")
    icon_path = resource_path("Public/Logo-NoBg-black.ico")
    progress_window.iconbitmap(icon_path)
    progress_window.resizable(False, False)

    progress_label = tk.Label(progress_window, text="Processing images...")
    progress_label.pack(pady=10)

    progress = ttk.Progressbar(progress_window, orient="horizontal",
                               length=250, mode="determinate", maximum=total_images)
    progress.pack(pady=10)

    # Make sure the main window is not interactive during progress
    progress_window.grab_set()

    return progress_window, progress


# Disable all buttons and entry fields in the main window
def disable_main_window(root):
    for widget in root.winfo_children():
        if isinstance(widget, (tk.Button, tk.Entry)):
            widget.config(state=tk.DISABLED)


# Enable all buttons and entry fields in the main window
def enable_main_window(root):
    for widget in root.winfo_children():
        if isinstance(widget, (tk.Button, tk.Entry)):
            widget.config(state=tk.NORMAL)


# GUI setup
def gui():
    root = tk.Tk()
    root.title("DokSofort - Bilder zu Word")
    root.geometry("500x320")
    root.resizable(width=False, height=False)
    root.configure(bg="black")
    icon_path = resource_path("Public/Logo-NoBg-black.ico")
    root.iconbitmap(icon_path)

    # Holds image directories
    image_folders = [""] * MAXIMUM_IMAGE_DIRS
    save_path_var = tk.StringVar()

    # Load the image
    image_path = resource_path("Public/Logo-NoBg.png")
    img = Image.open(image_path)

    # Resize the image
    resized_image = img.resize((80, 80), Image.Resampling.LANCZOS)
    image = ImageTk.PhotoImage(resized_image)

    # Frame to hold the title label and image
    title_frame = tk.Frame(root, bg="black")
    title_frame.pack(pady=10, anchor="w")

    # Image label
    image_label = tk.Label(title_frame, image=image, bg="black")
    image_label.pack(side=tk.LEFT)

    # Title label
    title_label = tk.Label(title_frame, text="Dokument aus Bildern Generieren", font=(
        'Arial', 16), anchor="w", fg="white", bg="black")
    title_label.pack(side=tk.LEFT, padx=10)

    # Main content frame, aligned to the left
    content_frame = tk.Frame(root, bg="black")
    content_frame.pack(anchor="w", padx=10, pady=10)

    # List to keep track of "Select Image Directory" frames
    directory_frames = []

    # Create add_more_button and remove_last_button
    button_frame = tk.Frame(content_frame, bg="black")
    button_frame.pack(anchor="w", pady=10)

    remove_last_button = tk.Button(
        button_frame, text="  -  ", font=('Arial', 12), state=tk.DISABLED, bg="gray", fg="white")
    add_more_button = tk.Button(
        button_frame, text=" + ", font=('Arial', 12), state=tk.DISABLED, bg="gray", fg="white")

    def update_remove_button_state():
        """Enable/disable the remove button based on the number of directory frames."""
        if len(directory_frames) > 1:
            remove_last_button.config(state=tk.NORMAL)
        else:
            remove_last_button.config(state=tk.DISABLED)

    def remove_last_image_directory_selector():
        """Remove the last 'Select Image Directory' frame."""
        if len(directory_frames) > 1:
            last_frame = directory_frames.pop()
            last_frame.destroy()

            # Remove the corresponding folder from the image_folders list
            image_folders[len(directory_frames)] = ""

            # Enable "Add More" button again
            if len(directory_frames) < 3:
                add_more_button.config(state=tk.NORMAL)

            # Decrease the window height when an image directory selector is removed
            current_height = root.winfo_height()
            root.geometry(f"500x{current_height - 40}")

            update_remove_button_state()

    def on_generate():
        # Filter out empty directories before generating the document
        valid_image_folders = [folder for folder in image_folders if folder]

        if not valid_image_folders:
            messagebox.showwarning("Warnung", "Kein Bilder-Ordner selektiert!")
            return

        if not save_path_var.get():
            messagebox.showwarning(
                "Warnung", "Kein Speicher-Ordner selektiert!")
            return

        # Disable main window widgets and show progress bar
        disable_main_window(root)
        total_images = sum(len(os.listdir(folder))
                           for folder in valid_image_folders if os.path.isdir(folder))
        progress_window, progress_bar = show_progress_bar(total_images)

        # Callback function to update progress
        def update_progress(current, total):
            progress_bar["value"] = current
            progress_bar.update()

        # Run the document generation in a separate thread to avoid blocking the GUI
        def create_document_thread():
            create_document(valid_image_folders,
                            save_path_var.get(), update_progress)
            progress_window.destroy()
            enable_main_window(root)

        threading.Thread(target=create_document_thread).start()

    def add_image_directory_selector():
        nonlocal add_more_button

        # Maximum Image dir
        if len(directory_frames) >= MAXIMUM_IMAGE_DIRS:
            return

        # Frame for directory selection
        frame = tk.Frame(content_frame, bg="black")
        frame.pack(anchor="w", pady=5, before=button_frame)

        # Button to select the image directory
        select_button = tk.Button(
            frame, text="Bilder-Ordner..", font=('Arial', 10), bg="gray", fg="white")
        select_button.pack(side=tk.LEFT, padx=10)

        # Label to show the selected directory
        label = tk.Label(frame, text="Kein Bilder-Ordner selektiert",
                         width=40, anchor="w", font=('Arial', 10), bg="black", fg="white")
        label.pack(side=tk.LEFT)

        # Bind the button to select folder and update state
        index = len(directory_frames)
        select_button.config(command=lambda: select_image_folder(
            label, image_folders, add_more_button, generate_button, save_path_var, index))

        # Add the frame to the list of directory frames
        directory_frames.append(frame)

        # Disable the "Add More" button if we reach the limit
        if len(directory_frames) >= MAXIMUM_IMAGE_DIRS:
            add_more_button.config(state=tk.DISABLED)
        else:
            add_more_button.config(state=tk.NORMAL)

        # Increase the window height when an image directory selector is added
        current_height = root.winfo_height()
        root.geometry(f"500x{current_height + 40}")

        # Enable/disable the remove button based on the number of directory frames
        update_remove_button_state()

    # Initially add one image directory selector
    add_image_directory_selector()

    # Configure the "-" button to remove the last selector
    remove_last_button.config(
        state=tk.DISABLED, command=remove_last_image_directory_selector)
    remove_last_button.pack(side=tk.LEFT, padx=10)

    # Now enable the "+" button after the first selector is added
    add_more_button.config(
        state=tk.NORMAL, command=add_image_directory_selector)
    add_more_button.pack(side=tk.LEFT, padx=0)

    # Frame for selecting output directory, aligned to the left
    output_frame = tk.Frame(content_frame, bg="black")
    output_frame.pack(anchor="w", pady=20)

    # Button to select the output directory
    output_button = tk.Button(output_frame, text="Output Ordner..", font=(
        'Arial', 10), bg="gray", fg="white")
    output_button.pack(side=tk.LEFT, padx=10)

    # Label to show the selected output directory
    output_label = tk.Label(output_frame, text="Kein Output Ordner selektiert",
                            width=40, anchor="w", font=('Arial', 10), bg="black", fg="white")
    output_label.pack(side=tk.LEFT)

    # Bind output directory selection
    output_button.config(command=lambda: select_save_location(
        output_label, save_path_var, image_folders, generate_button))

    # Generate button (initially disabled) with larger font and aligned to the left
    generate_button = tk.Button(content_frame, text="Dokument Generieren", command=on_generate, font=(
        'Arial', 12), state=tk.DISABLED, bg="#06402B", fg="white")
    generate_button.pack(anchor="center", pady=20)

    root.mainloop()


# Run the GUI
if __name__ == "__main__":
    gui()
